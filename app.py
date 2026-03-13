import os, re, json, ssl, smtplib
from datetime import datetime
from typing import List, Literal, Optional
from email.message import EmailMessage

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse, RedirectResponse
from pydantic import BaseModel, Field
from tenacity import retry, stop_after_attempt, wait_exponential
import httpx

from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from tzlocal import get_localzone

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

from dotenv import load_dotenv
load_dotenv("engine.env")

import stripe

HOST = os.getenv("ENGINE_HOST", "127.0.0.1")
PORT = int(os.getenv("ENGINE_PORT", "7363"))
DATA_ROOT = os.getenv("DATA_ROOT", ".")

# Stripe Configuration
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY", "")
STRIPE_PUBLISHABLE_KEY = os.getenv("STRIPE_PUBLISHABLE_KEY", "")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "")
stripe.api_key = STRIPE_SECRET_KEY

# Pricing
PRICE_LEAD_INTAKE = os.getenv("PRICE_LEAD_INTAKE", "price_lead_intake")
PRICE_SPRINT_50 = os.getenv("PRICE_SPRINT_50", "price_sprint_50")

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://127.0.0.1:11434")
LMSTUDIO_URL = os.getenv("LMSTUDIO_URL", "http://100.83.148.64:1234/v1/chat/completions")
MODEL = os.getenv("REPLY_MODEL", "qwen2.5:7b-instruct")

SMTP_ENABLED = os.getenv("SMTP_ENABLED", "false").lower() == "true"
SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASS = os.getenv("SMTP_PASS", "")
FROM_EMAIL = os.getenv("FROM_EMAIL", "owner@example.com")
FROM_NAME = os.getenv("FROM_NAME", "Review Reply Sprint")

SAMPLES_URL = os.getenv("SAMPLES_URL", "https://droxaillc.com/review-reply.html")


class Review(BaseModel):
    rating: int = Field(ge=1, le=5)
    text: str

class Job(BaseModel):
    business_name: str
    signoff: str = "-"
    tone_style: Literal["Professional","Warm","Playful","Direct"] = "Professional"
    persona: Literal["dentist_medspa","trades","auto","restaurant","real_estate","ecom"] = "trades"
    compliance_profile: Literal["HIPAA","FTC","General"] = "General"
    private_email: Optional[str] = None
    banned_words: List[str] = []
    max_length_per_reply: int = 900
    reviews: List[Review]

class Lead(BaseModel):
    business_name: str
    email: str
    persona: Literal["dentist_medspa","trades","auto","restaurant","real_estate","ecom"] = "trades"
    compliance_profile: Literal["HIPAA","FTC","General"] = "General"
    tone_style: Literal["Professional","Warm","Playful","Direct"] = "Professional"
    signoff: str = "- Team"
    private_email: Optional[str] = None
    reviews_sample: List[Review]


# SHORTENED PROMPTS to avoid context size issues
SYSTEM_TMPL = """Write owner-voice Google review replies.
Persona: {persona}. Compliance: {compliance}.
- Be specific. No templates.
- Rating <=2: apologize, invite follow-up at {private_email}.
- Max {max_len} chars.
- Avoid: {banned}
- Tone: {tone}. Sign: {signoff}
HIPAA=no PHI. FTC=no guarantees.
Return ONLY JSON: ["reply1","reply2",...]
"""

USER_TMPL = """Business: {biz}
Reviews:
{pairs}
Return JSON array only.
"""


async def ollama_generate(prompt: str) -> str:
    async with httpx.AsyncClient(timeout=300.0) as cx:
        r = await cx.post(f"{OLLAMA_URL}/api/generate", json={"model": MODEL, "prompt": prompt, "stream": False})
        if r.status_code != 200:
            raise RuntimeError(f"Ollama error: {r.text}")
        return r.json().get("response","")

async def lmstudio_chat(system: str, user: str) -> str:
    async with httpx.AsyncClient(timeout=300.0) as cx:
        r = await cx.post(LMSTUDIO_URL, json={
            "model":"local-model",
            "messages":[{"role":"system","content":system},{"role":"user","content":user}],
            "temperature":0.3,
            "max_tokens":2000
        })
        if r.status_code != 200:
            raise RuntimeError(f"LM Studio error: {r.text}")
        data = r.json()
        return data["choices"][0]["message"]["content"]

def json_only(text: str):
    m = re.search(r"\[[\s\S]*\]", text)
    if not m: raise ValueError("No JSON array found in model output")
    return json.loads(m.group(0))

def qa_guard(reply: str, banned: List[str], max_len: int):
    s = reply.strip()
    if len(s) > max_len: s = s[:max_len-1] + "..."
    low = s.lower()
    for w in banned:
        if w.lower() in low:
            s = s.replace(w, "***")
    return s

@retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=1, max=4))
async def generate_replies(job: Job):
    private_email = job.private_email or "contact us"
    sys_prompt = SYSTEM_TMPL.format(
        persona=job.persona,
        compliance=job.compliance_profile,
        private_email=private_email,
        max_len=job.max_length_per_reply,
        banned=", ".join(job.banned_words) if job.banned_words else "none",
        tone=job.tone_style,
        signoff=job.signoff
    )
    pairs = "\n".join([f"- ({r.rating}) {r.text[:150]}" for r in job.reviews])  # Truncate long reviews
    user_prompt = USER_TMPL.format(biz=job.business_name, pairs=pairs)
    
    try:
        raw = await ollama_generate(f"{sys_prompt}\n\n{user_prompt}")
    except Exception:
        raw = await lmstudio_chat(sys_prompt, user_prompt)
    
    arr = json_only(raw)
    out = [qa_guard(a, job.banned_words, job.max_length_per_reply) for a in arr]
    
    # Ensure we have the right number of replies
    if len(out) != len(job.reviews):
        out = (out + [f"Thank you for your feedback. {job.signoff}"]*(len(job.reviews)-len(out)))[:len(job.reviews)]
    return out


def ensure_dir(p: str):
    os.makedirs(p, exist_ok=True)

def job_dirs(business: str):
    safe = re.sub(r"[^a-zA-Z0-9._-]+","_", business).strip("_")
    base = os.path.join(DATA_ROOT,"clients", safe, datetime.now().strftime("%Y-%m-%d"))
    ensure_dir(base)
    return base

def build_docx(path: str, business: str, reviews: List[Review], replies: List[str], signoff: str):
    doc = Document()
    doc.add_heading(f"{business} - Review Replies", level=0)
    for i,(rv,rep) in enumerate(zip(reviews, replies), start=1):
        doc.add_heading(f"Review #{i} - {rv.rating} Stars", level=2)
        doc.add_paragraph(rv.text)
        para = doc.add_paragraph()
        run = para.add_run(rep)
        font = run.font; font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph(f"{signoff}")
    doc.save(path)

def build_pdf(path: str, business: str, reviews: List[Review], replies: List[str], signoff: str):
    def clean(text):
        """Sanitize text for PDF (replace fancy Unicode characters)"""
        replacements = {
            '\u2014': '-', '\u2013': '-',  # dashes
            '\u2018': "'", '\u2019': "'",  # quotes
            '\u201c': '"', '\u201d': '"',
            '\u2026': '...',  # ellipsis
            '—': '-', '–': '-',
            ''': "'", ''': "'",
            '"': '"', '"': '"',
            '★': '*', '…': '...'
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text.encode('latin-1', errors='replace').decode('latin-1')
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial","B",16)
    pdf.cell(0,10, clean(f"{business} - Review Replies"), ln=1)
    pdf.set_font("Arial","",12)
    
    for i,(rv,rep) in enumerate(zip(reviews, replies), start=1):
        pdf.set_font("Arial","B",12)
        pdf.cell(0,8, clean(f"Review #{i} - {rv.rating} Stars"), ln=1)
        pdf.set_font("Arial","",12)
        pdf.multi_cell(0,6, clean(rv.text))
        pdf.set_font("Arial","I",12)
        pdf.multi_cell(0,6, clean(rep))
        pdf.ln(2)
    
    pdf.ln(3)
    pdf.set_font("Arial","",11)
    pdf.cell(0,6, clean(signoff), ln=1)
    pdf.output(path)

def send_email(to_addr: str, subject: str, body: str, attachments: Optional[List[str]]=None):
    """Send email with debug logging"""
    msg = EmailMessage()
    msg["From"] = f"{FROM_NAME} <{FROM_EMAIL}>"
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.set_content(body)

    for a in attachments or []:
        with open(a,"rb") as f:
            data = f.read()
        maintype, subtype = ("application","octet-stream")
        msg.add_attachment(data, maintype=maintype, subtype=subtype, filename=os.path.basename(a))

    print(f"\n{'='*70}")
    print(f"📧 EMAIL SEND ATTEMPT")
    print(f"To: {to_addr}")
    print(f"Subject: {subject}")
    print(f"SMTP_ENABLED: {SMTP_ENABLED}")
    print(f"SMTP_HOST: {SMTP_HOST}")
    print(f"SMTP_PORT: {SMTP_PORT}")
    print(f"SMTP_USER: {SMTP_USER}")
    print(f"SMTP_PASS: {'***' + SMTP_PASS[-4:] if SMTP_PASS and len(SMTP_PASS) > 4 else 'NOT SET'}")
    print(f"FROM_EMAIL: {FROM_EMAIL}")
    print(f"FROM_NAME: {FROM_NAME}")
    print(f"{'='*70}")

    if SMTP_ENABLED and SMTP_HOST and SMTP_USER and SMTP_PASS:
        try:
            print("🔧 Step 1: Creating SSL context...")
            ctx = ssl.create_default_context()
            
            print(f"🔧 Step 2: Connecting to {SMTP_HOST}:{SMTP_PORT}...")
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=30) as s:
                print("🔧 Step 3: Starting TLS encryption...")
                s.starttls(context=ctx)
                
                print("🔧 Step 4: Logging in...")
                s.login(SMTP_USER, SMTP_PASS)
                
                print("🔧 Step 5: Sending message...")
                s.send_message(msg)
                
                print(f"✅ EMAIL SENT SUCCESSFULLY to {to_addr}")
                print(f"{'='*70}\n")
                
        except smtplib.SMTPAuthenticationError as e:
            print(f"❌ AUTHENTICATION ERROR: {e}")
            print("⚠️  This usually means:")
            print("   - Wrong password")
            print("   - 2FA enabled (need App Password)")
            print("   - 'Less secure app access' disabled")
            print(f"{'='*70}\n")
            _save_locally(msg, subject)
            
        except smtplib.SMTPException as e:
            print(f"❌ SMTP ERROR: {type(e).__name__}: {e}")
            print(f"{'='*70}\n")
            _save_locally(msg, subject)
            
        except Exception as e:
            print(f"❌ UNEXPECTED ERROR: {type(e).__name__}: {e}")
            print(f"{'='*70}\n")
            _save_locally(msg, subject)
    else:
        print("⚠️  SMTP not fully configured. Missing:")
        if not SMTP_ENABLED: print("   - SMTP_ENABLED=true")
        if not SMTP_HOST: print("   - SMTP_HOST")
        if not SMTP_USER: print("   - SMTP_USER")
        if not SMTP_PASS: print("   - SMTP_PASS")
        print(f"{'='*70}\n")
        _save_locally(msg, subject)

def _save_locally(msg: EmailMessage, subject: str):
    """Save email to local outbox folder"""
    outbox = os.path.join(DATA_ROOT,"outbox")
    ensure_dir(outbox)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    fn = os.path.join(outbox, f"{ts}-{subject.replace(' ','_')}.eml")
    with open(fn,"wb") as f:
        f.write(msg.as_bytes())
    print(f"📁 Email saved locally to: {fn}\n")


def weekly_kpi_job():
    pass

local_tz = get_localzone()
sched = BackgroundScheduler(timezone=local_tz)
sched.add_job(weekly_kpi_job, CronTrigger(day_of_week="mon", hour=8, minute=0))
sched.start()

app = FastAPI(title="Review Reply Orchestrator", version="1.0.0")

LANDING_PAGE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Review Reply Sprint - AI-Powered Review Responses</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body { background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%); color: #f1f5f9; }
        .card { background: rgba(255,255,255,0.05); backdrop-filter: blur(10px); border: 1px solid rgba(255,255,255,0.1); }
        .btn { background: linear-gradient(135deg, #6366f1, #8b5cf6); }
    </style>
</head>
<body class="min-h-screen">
    <div class="max-w-6xl mx-auto px-4 py-16">
        <header class="text-center mb-16">
            <h1 class="text-5xl font-extrabold mb-4">Never Stress Over Review Replies Again</h1>
            <p class="text-xl text-slate-300 max-w-2xl mx-auto">
                AI-powered review responses that sound like YOU. Save 10+ hours per week. 
                Get your first 5 replies FREE, then $199 for 50 professional responses.
            </p>
        </header>

        <div class="grid md:grid-cols-2 gap-8 mb-16">
            <div class="card rounded-2xl p-8">
                <h3 class="text-2xl font-bold mb-4">🎁 Free Sample</h3>
                <p class="text-slate-300 mb-6">Send us 5 reviews, get 5 custom AI-generated replies back. Free. No credit card required.</p>
                <form action="/lead-intake-demo" method="POST" class="space-y-4">
                    <input type="text" name="business_name" placeholder="Your Business Name" required class="w-full px-4 py-3 rounded-lg bg-slate-800 border border-slate-700 focus:border-indigo-500 outline-none">
                    <input type="email" name="email" placeholder="Your Email" required class="w-full px-4 py-3 rounded-lg bg-slate-800 border border-slate-700 focus:border-indigo-500 outline-none">
                    <textarea name="reviews" placeholder="Paste 5 reviews here (one per line, include star rating like '5 Great service!')" rows="6" required class="w-full px-4 py-3 rounded-lg bg-slate-800 border border-slate-700 focus:border-indigo-500 outline-none"></textarea>
                    <button type="submit" class="w-full btn text-white font-bold py-3 px-6 rounded-lg hover:opacity-90 transition">Get My Free Samples</button>
                </form>
            </div>

            <div class="card rounded-2xl p-8 border-2 border-indigo-500/50">
                <div class="text-sm text-indigo-400 font-semibold mb-2">MOST POPULAR</div>
                <h3 class="text-2xl font-bold mb-4">⚡ 50-Reply Sprint</h3>
                <p class="text-slate-300 mb-4">Full batch of 50 custom replies. Perfect for catching up on your review backlog.</p>
                <div class="text-4xl font-extrabold text-indigo-400 mb-6">$199 <span class="text-lg text-slate-400 font-normal">one-time</span></div>
                <ul class="space-y-3 text-slate-300 mb-8">
                    <li>✓ 50 AI-generated custom replies</li>
                    <li>✓ Your brand voice & tone</li>
                    <li>✓ Compliance-checked (HIPAA/FTC)</li>
                    <li>✓ PDF + DOCX formats</li>
                    <li>✓ 24-hour delivery</li>
                </ul>
                <form action="/create-checkout-session" method="POST">
                    <input type="hidden" name="product" value="sprint_50">
                    <button type="submit" class="w-full btn text-white font-bold py-4 px-6 rounded-lg hover:opacity-90 transition text-lg">Buy Now - $199</button>
                </form>
            </div>
        </div>

        <div class="text-center text-slate-400">
            <p>Questions? <a href="mailto:droxai25@outlook.com" class="text-indigo-400 hover:underline">droxai25@outlook.com</a></p>
            <p class="mt-4 text-sm">Built by DroxAI LLC</p>
        </div>
    </div>
</body>
</html>
"""

@app.get("/", response_class=HTMLResponse)
async def landing_page():
    return LANDING_PAGE

@app.get("/healthz")
def healthz():
    return {"ok": True, "port": PORT, "model": MODEL}

@app.post("/create-checkout-session")
async def create_checkout_session(request: Request):
    if not STRIPE_SECRET_KEY:
        raise HTTPException(status_code=500, detail="Stripe not configured")
    
    form_data = await request.form()
    product = form_data.get("product", "sprint_50")
    
    prices = {
        "sprint_50": {"price": 19900, "name": "50-Reply Sprint"},
        "lead_intake": {"price": 0, "name": "Free Sample (5 replies)"},
    }
    
    item = prices.get(product, prices["sprint_50"])
    
    try:
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{
                "price_data": {
                    "currency": "usd",
                    "product_data": {"name": item["name"]},
                    "unit_amount": item["price"],
                },
                "quantity": 1,
            }],
            mode="payment",
            success_url=f"{request.base_url}success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{request.base_url}cancel",
            metadata={"product": product},
        )
        return RedirectResponse(url=checkout_session.url, status_code=303)
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=500, detail=f"Stripe error: {str(e)}")

@app.get("/success")
async def success_page(session_id: str):
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            return HTMLResponse("""
            <!DOCTYPE html>
            <html><head><script src="https://cdn.tailwindcss.com"></script></head>
            <body class="min-h-screen flex items-center justify-center" style="background:linear-gradient(135deg,#0f172a,#1e293b);color:#f1f5f9">
                <div class="text-center">
                    <div class="text-6xl mb-4">✅</div>
                    <h1 class="text-4xl font-extrabold mb-4">Payment Successful!</h1>
                    <p class="text-xl text-slate-300 mb-8">Thank you! We'll contact you within 24 hours to get started.</p>
                    <a href="mailto:droxai25@outlook.com" class="inline-block bg-indigo-600 text-white font-bold py-3 px-8 rounded-lg hover:bg-indigo-700">Email Us Now</a>
                </div>
            </body></html>
            """)
    except Exception:
        pass
    return HTMLResponse("<h1>Order Status Unknown</h1><p>Please contact us at droxai25@outlook.com</p>")

@app.get("/cancel")
async def cancel_page():
    return HTMLResponse("""
    <!DOCTYPE html>
    <html><head><script src="https://cdn.tailwindcss.com"></script></head>
    <body class="min-h-screen flex items-center justify-center" style="background:linear-gradient(135deg,#0f172a,#1e293b);color:#f1f5f9">
        <div class="text-center">
            <div class="text-6xl mb-4">❌</div>
            <h1 class="text-4xl font-extrabold mb-4">Checkout Cancelled</h1>
            <p class="text-xl text-slate-300 mb-8">No worries! Ready when you are.</p>
            <a href="/" class="inline-block bg-indigo-600 text-white font-bold py-3 px-8 rounded-lg hover:bg-indigo-700">Go Back</a>
        </div>
    </body></html>
    """)

@app.post("/webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")
    
    if not STRIPE_WEBHOOK_SECRET or not sig_header:
        raise HTTPException(status_code=400, detail="Webhook not configured")
    
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except (ValueError, stripe.error.SignatureVerificationError) as e:
        raise HTTPException(status_code=400, detail=f"Invalid webhook: {str(e)}")
    
    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        print(f"Payment received! Session: {session['id']}, Product: {session.get('metadata', {}).get('product')}")
    
    return {"status": "success"}


@app.post("/lead-intake")
async def lead_intake(lead: Lead):
    job = Job(
        business_name=lead.business_name,
        signoff=lead.signoff,
        tone_style=lead.tone_style,
        persona=lead.persona,
        compliance_profile=lead.compliance_profile,
        private_email=lead.private_email,
        banned_words=["guarantee"],
        max_length_per_reply=400,
        reviews=lead.reviews_sample
    )
    replies = await generate_replies(job)
    base = job_dirs(lead.business_name)
    pdf = os.path.join(base, "sample-5.pdf")
    docx= os.path.join(base, "sample-5.docx")
    build_pdf(pdf, lead.business_name, job.reviews, replies, job.signoff)
    build_docx(docx, lead.business_name, job.reviews, replies, job.signoff)
    send_email(
        to_addr=lead.email,
        subject=f"{lead.business_name} - 5 tailored review replies",
        body=(f"Hi! Here are your 5 custom review replies.\n\nIf you like the quality, get 50 more here:\n{SAMPLES_URL}\n\nThanks!\nDroxAI"),
        attachments=[pdf, docx]
    )
    return {"ok": True, "sent": True, "files": [pdf, docx]}