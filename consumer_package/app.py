import asyncio
import logging
import os
import re
import json
from datetime import datetime
from typing import List, Literal, Optional

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse, RedirectResponse
from pydantic import BaseModel, Field
from tenacity import retry, stop_after_attempt, wait_exponential
import httpx

from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from tzlocal import get_localzone

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

from dotenv import load_dotenv

load_dotenv("engine.env")

import stripe

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)


def check_ollama_running():
    """Check if Ollama is running and accessible."""
    import httpx
    try:
        response = httpx.get(OLLAMA_URL, timeout=5.0)
        return response.status_code == 200
    except Exception:
        return False

HOST = os.getenv("ENGINE_HOST", "127.0.0.1")
PORT = int(os.getenv("ENGINE_PORT", "7363"))
DATA_ROOT = os.getenv("DATA_ROOT", ".")

# Stripe Configuration
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY", "")
STRIPE_PUBLISHABLE_KEY = os.getenv("STRIPE_PUBLISHABLE_KEY", "")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "")
stripe.api_key = STRIPE_SECRET_KEY

# Pricing
PRICE_LEAD_INTAKE = os.getenv("PRICE_LEAD_INTAKE", "price_lead_intake")
PRICE_SPRINT_50 = os.getenv("PRICE_SPRINT_50", "price_sprint_50")

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://127.0.0.1:11434")
LMSTUDIO_URL = os.getenv(
    "LMSTUDIO_URL", "http://100.83.148.64:1234/v1/chat/completions"
)
MODEL = os.getenv("REPLY_MODEL", "qwen2.5:7b-instruct")

SAMPLES_URL = os.getenv("SAMPLES_URL", "https://droxaillc.com/review-reply.html")


class Review(BaseModel):
    rating: int = Field(ge=1, le=5)
    text: str


class Job(BaseModel):
    business_name: str
    signoff: str = "-"
    tone_style: Literal["Professional", "Warm", "Playful", "Direct"] = "Professional"
    persona: Literal[
        "dentist_medspa", "trades", "auto", "restaurant", "real_estate", "ecom"
    ] = "trades"
    compliance_profile: Literal["HIPAA", "FTC", "General"] = "General"
    private_email: Optional[str] = None
    banned_words: List[str] = []
    max_length_per_reply: int = 900
    reviews: List[Review]


class Lead(BaseModel):
    business_name: str
    email: str
    persona: Literal[
        "dentist_medspa", "trades", "auto", "restaurant", "real_estate", "ecom"
    ] = "trades"
    compliance_profile: Literal["HIPAA", "FTC", "General"] = "General"
    tone_style: Literal["Professional", "Warm", "Playful", "Direct"] = "Professional"
    signoff: str = "- Team"
    private_email: Optional[str] = None
    reviews_sample: List[Review]


class AutomationCredentials(BaseModel):
    email: str
    password: str


class AutomationConfig(BaseModel):
    enabled: Optional[bool] = None
    interval_minutes: Optional[int] = None
    headless: Optional[bool] = None
    business_name: Optional[str] = None


# SHORTENED PROMPTS to avoid context size issues
SYSTEM_TMPL = """Write owner-voice Google review replies.
Persona: {persona}. Compliance: {compliance}.
- Be specific. No templates.
- Rating <=2: apologize, invite follow-up at {private_email}.
- Max {max_len} chars.
- Avoid: {banned}
- Tone: {tone}. Sign: {signoff}
HIPAA=no PHI. FTC=no guarantees.
Return ONLY JSON: ["reply1","reply2",...]
"""

USER_TMPL = """Business: {biz}
Reviews:
{pairs}
Return JSON array only.
"""


async def ollama_generate(prompt: str) -> str:
    async with httpx.AsyncClient(timeout=300.0) as cx:
        r = await cx.post(
            f"{OLLAMA_URL}/api/generate",
            json={"model": MODEL, "prompt": prompt, "stream": False},
        )
        if r.status_code != 200:
            raise RuntimeError(f"Ollama error: {r.text}")
        return r.json().get("response", "")


async def lmstudio_chat(system: str, user: str) -> str:
    async with httpx.AsyncClient(timeout=300.0) as cx:
        r = await cx.post(
            LMSTUDIO_URL,
            json={
                "model": "local-model",
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                "temperature": 0.3,
                "max_tokens": 2000,
            },
        )
        if r.status_code != 200:
            raise RuntimeError(f"LM Studio error: {r.text}")
        data = r.json()
        return data["choices"][0]["message"]["content"]


def json_only(text: str):
    m = re.search(r"\[[\s\S]*\]", text)
    if not m:
        raise ValueError("No JSON array found in model output")
    return json.loads(m.group(0))


def qa_guard(reply: str, banned: List[str], max_len: int):
    s = reply.strip()
    if len(s) > max_len:
        s = s[: max_len - 1] + "..."
    low = s.lower()
    for w in banned:
        if w.lower() in low:
            s = s.replace(w, "***")
    return s


@retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=1, max=4))
async def generate_replies(job: Job):
    private_email = job.private_email or "contact us"
    sys_prompt = SYSTEM_TMPL.format(
        persona=job.persona,
        compliance=job.compliance_profile,
        private_email=private_email,
        max_len=job.max_length_per_reply,
        banned=", ".join(job.banned_words) if job.banned_words else "none",
        tone=job.tone_style,
        signoff=job.signoff,
    )
    pairs = "\n".join(
        [f"- ({r.rating}) {r.text[:150]}" for r in job.reviews]
    )  # Truncate long reviews
    user_prompt = USER_TMPL.format(biz=job.business_name, pairs=pairs)

    try:
        raw = await ollama_generate(f"{sys_prompt}\n\n{user_prompt}")
    except Exception:
        raw = await lmstudio_chat(sys_prompt, user_prompt)

    arr = json_only(raw)
    out = [qa_guard(a, job.banned_words, job.max_length_per_reply) for a in arr]

    # Ensure we have the right number of replies
    if len(out) != len(job.reviews):
        out = (
            out
            + [f"Thank you for your feedback. {job.signoff}"]
            * (len(job.reviews) - len(out))
        )[: len(job.reviews)]
    return out


def ensure_dir(p: str):
    os.makedirs(p, exist_ok=True)


def job_dirs(business: str):
    safe = re.sub(r"[^a-zA-Z0-9._-]+", "_", business).strip("_")
    base = os.path.join(DATA_ROOT, "clients", safe, datetime.now().strftime("%Y-%m-%d"))
    ensure_dir(base)
    return base


def build_docx(
    path: str, business: str, reviews: List[Review], replies: List[str], signoff: str
):
    doc = Document()
    doc.add_heading(f"{business} - Review Replies", level=0)
    for i, (rv, rep) in enumerate(zip(reviews, replies), start=1):
        doc.add_heading(f"Review #{i} - {rv.rating} Stars", level=2)
        doc.add_paragraph(rv.text)
        para = doc.add_paragraph()
        run = para.add_run(rep)
        font = run.font
        font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph(f"{signoff}")
    doc.save(path)


def build_pdf(
    path: str, business: str, reviews: List[Review], replies: List[str], signoff: str
):
    def clean(text):
        """Sanitize text for PDF (replace fancy Unicode characters)"""
        replacements = {
            "\u2014": "-",
            "\u2013": "-",  # dashes
            "\u2018": "'",
            "\u2019": "'",  # quotes
            "\u201c": '"',
            "\u201d": '"',
            "\u2026": "...",  # ellipsis
            "—": "-",
            "–": "-",
            """: "'", """: "'",
            '"': '"',
            '"': '"',
            "★": "*",
            "…": "...",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, clean(f"{business} - Review Replies"), ln=1)
    pdf.set_font("Arial", "", 12)

    for i, (rv, rep) in enumerate(zip(reviews, replies), start=1):
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, clean(f"Review #{i} - {rv.rating} Stars"), ln=1)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(0, 6, clean(rv.text))
        pdf.set_font("Arial", "I", 12)
        pdf.multi_cell(0, 6, clean(rep))
        pdf.ln(2)

    pdf.ln(3)
    pdf.set_font("Arial", "", 11)
    pdf.cell(0, 6, clean(signoff), ln=1)
    pdf.output(path)


local_tz = get_localzone()
sched = BackgroundScheduler(timezone=local_tz)
sched.start()

app = FastAPI(title="Review Reply Orchestrator", version="2.0.0")


@app.on_event("startup")
async def startup_event():
    """Start background review monitor and automation service on app startup."""
    from monitor import monitor_loop  # noqa: PLC0415

    asyncio.create_task(monitor_loop())
    logger.info("Review monitoring loop started.")

    # Auto-start browser automation if credentials are configured
    try:
        from browser_automation import credentials_configured  # noqa: PLC0415
        from automation_service import start_service  # noqa: PLC0415

        if credentials_configured():
            start_service()
            logger.info("Browser automation service auto-started.")
        else:
            logger.info(
                "Browser automation credentials not configured. "
                "Visit /automation/setup to set them up."
            )
    except Exception as exc:
        logger.warning("Could not start browser automation service: %s", exc)


LANDING_PAGE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Review Reply Orchestrator - Automatic Google Review Replies</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body { background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%); color: #f1f5f9; }
        .card { background: rgba(255,255,255,0.05); backdrop-filter: blur(10px); border: 1px solid rgba(255,255,255,0.1); }
        .btn { background: linear-gradient(135deg, #6366f1, #8b5cf6); }
        .btn-green { background: linear-gradient(135deg, #16a34a, #15803d); }
        .btn-red { background: linear-gradient(135deg, #dc2626, #b91c1c); }
    </style>
</head>
<body class="min-h-screen">
    <div class="max-w-6xl mx-auto px-4 py-16">
        <header class="text-center mb-16">
            <h1 class="text-5xl font-extrabold mb-4">Never Stress Over Review Replies Again</h1>
            <p class="text-xl text-slate-300 max-w-2xl mx-auto">
                AI-powered review responses posted automatically to Google Business.
                No API keys required. Just your Google login credentials.
            </p>
        </header>

        <!-- Browser Automation Setup -->
        <div class="card rounded-2xl p-8 mb-8 border-2 border-green-500/50">
            <h3 class="text-2xl font-bold mb-2">🤖 Automatic Review Replies (Browser Automation)</h3>
            <p class="text-slate-300 mb-6">
                Enter your Google Business credentials once. The software logs in automatically,
                checks for new reviews every hour, and posts AI-generated replies — no manual work.
            </p>
            <div class="grid md:grid-cols-2 gap-8">
                <div>
                    <h4 class="font-semibold mb-3 text-green-400">Step 1 – Enter Credentials</h4>
                    <form id="setup-form" onsubmit="setupAutomation(event)">
                        <input type="email" id="setup-email" placeholder="Google Business Email"
                            class="w-full bg-slate-800 border border-slate-600 rounded-lg px-4 py-3 mb-3 text-white placeholder-slate-400 focus:outline-none focus:border-green-500">
                        <input type="password" id="setup-password" placeholder="Google Business Password"
                            class="w-full bg-slate-800 border border-slate-600 rounded-lg px-4 py-3 mb-3 text-white placeholder-slate-400 focus:outline-none focus:border-green-500">
                        <button type="submit" class="w-full btn-green text-white font-bold py-3 px-6 rounded-lg hover:opacity-90 transition">
                            Save &amp; Start Automation
                        </button>
                    </form>
                    <p id="setup-msg" class="mt-3 text-sm"></p>
                </div>
                <div>
                    <h4 class="font-semibold mb-3 text-indigo-400">Step 2 – Control Service</h4>
                    <div id="auto-status" class="text-slate-400 mb-4 text-sm">Loading status…</div>
                    <div class="flex gap-3 flex-wrap">
                        <button onclick="controlService('start')" class="btn-green text-white font-bold py-2 px-5 rounded-lg hover:opacity-90 text-sm">▶ Start</button>
                        <button onclick="controlService('stop')" class="btn-red text-white font-bold py-2 px-5 rounded-lg hover:opacity-90 text-sm">■ Stop</button>
                        <button onclick="runNow()" class="btn text-white font-bold py-2 px-5 rounded-lg hover:opacity-90 text-sm">⚡ Run Now</button>
                        <a href="/automation/logs" class="btn text-white font-bold py-2 px-5 rounded-lg hover:opacity-90 text-sm inline-block">📋 View Logs</a>
                    </div>
                    <p id="service-msg" class="mt-3 text-sm"></p>
                </div>
            </div>
        </div>

        <!-- Configure Interval -->
        <div class="card rounded-2xl p-6 mb-8">
            <h3 class="text-lg font-bold mb-3">⚙️ Configure Automation</h3>
            <div class="flex gap-4 items-center flex-wrap">
                <label class="text-slate-300">Check every</label>
                <input type="number" id="interval-input" value="60" min="1" max="1440"
                    class="w-24 bg-slate-800 border border-slate-600 rounded-lg px-3 py-2 text-white focus:outline-none focus:border-indigo-500">
                <label class="text-slate-300">minutes</label>
                <label class="flex items-center gap-2 text-slate-300">
                    <input type="checkbox" id="headless-check" checked class="w-4 h-4">
                    Headless (background)
                </label>
                <button onclick="saveConfig()" class="btn text-white font-bold py-2 px-5 rounded-lg hover:opacity-90 text-sm">Save Config</button>
                <span id="config-msg" class="text-sm text-green-400"></span>
            </div>
        </div>

        <div class="grid md:grid-cols-2 gap-8 mb-16">
            <div class="card rounded-2xl p-8 border-2 border-slate-500/50">
                <h3 class="text-2xl font-bold mb-4">🔗 Google API Connection (Optional)</h3>
                <p class="text-slate-300 mb-6">
                    Alternative: connect via Google OAuth API (requires Google Cloud setup).
                </p>
                <a href="/oauth/start" class="block w-full btn text-white font-bold py-3 px-6 rounded-lg hover:opacity-90 transition text-center">
                    Connect via Google API
                </a>
            </div>

            <div class="card rounded-2xl p-8 border-2 border-indigo-500/50">
                <div class="text-sm text-indigo-400 font-semibold mb-2">MOST POPULAR</div>
                <h3 class="text-2xl font-bold mb-4">⚡ 50-Reply Sprint</h3>
                <p class="text-slate-300 mb-4">Full batch of 50 custom replies. Perfect for catching up on your review backlog.</p>
                <div class="text-4xl font-extrabold text-indigo-400 mb-6">$199 <span class="text-lg text-slate-400 font-normal">one-time</span></div>
                <ul class="space-y-3 text-slate-300 mb-8">
                    <li>✓ 50 AI-generated custom replies</li>
                    <li>✓ Your brand voice &amp; tone</li>
                    <li>✓ Compliance-checked (HIPAA/FTC)</li>
                    <li>✓ PDF + DOCX formats</li>
                    <li>✓ 24-hour delivery</li>
                </ul>
                <form action="/create-checkout-session" method="POST">
                    <input type="hidden" name="product" value="sprint_50">
                    <button type="submit" class="w-full btn text-white font-bold py-4 px-6 rounded-lg hover:opacity-90 transition text-lg">Buy Now - $199</button>
                </form>
            </div>
        </div>

        <div class="card rounded-2xl p-8 mb-16">
            <h3 class="text-2xl font-bold mb-4">📊 Monitoring Status</h3>
            <div id="status-placeholder" class="text-slate-300">
                <a href="/reviews/status" class="text-indigo-400 hover:underline">Check API status →</a>
                &nbsp;|&nbsp;
                <a href="/automation/status" class="text-green-400 hover:underline">Check automation status →</a>
                &nbsp;|&nbsp;
                <a href="#" class="text-yellow-400 hover:underline" onclick="this.innerText='Syncing...'; fetch('/reviews/sync',{method:'POST'}).then(r=>r.json()).then(d=>{this.innerText='Sync complete ✓'}); return false;">
                    Trigger manual API sync →
                </a>
            </div>
        </div>

        <div class="text-center text-slate-400">
            <p>Questions? <a href="mailto:droxai25@outlook.com" class="text-indigo-400 hover:underline">droxai25@outlook.com</a></p>
            <p class="mt-4 text-sm">Built by DroxAI LLC</p>
        </div>
    </div>

    <script>
    // Load automation status on page load
    async function loadStatus() {
        try {
            const r = await fetch('/automation/status');
            const d = await r.json();
            const el = document.getElementById('auto-status');
            const running = d.running ? '🟢 Running' : '🔴 Stopped';
            const creds = d.credentials_configured ? '✅ Credentials saved' : '⚠️ No credentials';
            const last = d.last_run ? new Date(d.last_run).toLocaleString() : 'Never';
            el.innerHTML = `${running} &nbsp;|&nbsp; ${creds} &nbsp;|&nbsp; Last run: ${last} &nbsp;|&nbsp; Interval: ${d.interval_minutes}min`;

            // Pre-fill interval
            document.getElementById('interval-input').value = d.interval_minutes || 60;
        } catch(e) {
            document.getElementById('auto-status').textContent = 'Could not load status';
        }
    }

    async function setupAutomation(e) {
        e.preventDefault();
        const email = document.getElementById('setup-email').value;
        const password = document.getElementById('setup-password').value;
        const msg = document.getElementById('setup-msg');
        if (!email || !password) { msg.textContent = 'Email and password are required.'; msg.className='mt-3 text-sm text-red-400'; return; }
        msg.textContent = 'Saving…'; msg.className='mt-3 text-sm text-slate-400';
        try {
            const r = await fetch('/automation/setup', {
                method: 'POST',
                headers: {'Content-Type':'application/json'},
                body: JSON.stringify({email, password})
            });
            const d = await r.json();
            msg.textContent = d.ok ? '✅ Credentials saved & automation started!' : ('❌ ' + (d.detail || d.message));
            msg.className = 'mt-3 text-sm ' + (d.ok ? 'text-green-400' : 'text-red-400');
            document.getElementById('setup-password').value = '';
            if (d.ok) loadStatus();
        } catch(e) { msg.textContent = '❌ Request failed'; msg.className='mt-3 text-sm text-red-400'; }
    }

    async function controlService(action) {
        const msg = document.getElementById('service-msg');
        msg.textContent = 'Working…'; msg.className='mt-3 text-sm text-slate-400';
        try {
            const r = await fetch('/automation/' + action, {method:'POST'});
            const d = await r.json();
            msg.textContent = (d.ok ? '✅ ' : '⚠️ ') + d.message;
            msg.className = 'mt-3 text-sm ' + (d.ok ? 'text-green-400' : 'text-yellow-400');
            loadStatus();
        } catch(e) { msg.textContent = '❌ Request failed'; msg.className='mt-3 text-sm text-red-400'; }
    }

    async function runNow() {
        const msg = document.getElementById('service-msg');
        msg.textContent = 'Running automation pass… (may take a minute)'; msg.className='mt-3 text-sm text-slate-400';
        try {
            const r = await fetch('/automation/run-now', {method:'POST'});
            const d = await r.json();
            if (d.ok) {
                msg.textContent = `✅ Done – found ${d.reviews_found} reviews, posted ${d.replies_posted} replies, ${d.errors} errors`;
                msg.className='mt-3 text-sm text-green-400';
            } else {
                msg.textContent = '❌ ' + (d.error || 'Failed');
                msg.className='mt-3 text-sm text-red-400';
            }
        } catch(e) { msg.textContent = '❌ Request failed'; msg.className='mt-3 text-sm text-red-400'; }
    }

    async function saveConfig() {
        const interval = parseInt(document.getElementById('interval-input').value);
        const headless = document.getElementById('headless-check').checked;
        const msg = document.getElementById('config-msg');
        try {
            const r = await fetch('/automation/config', {
                method:'POST',
                headers:{'Content-Type':'application/json'},
                body: JSON.stringify({interval_minutes: interval, headless})
            });
            const d = await r.json();
            msg.textContent = d.ok ? '✅ Saved' : '❌ Failed';
            setTimeout(()=>{msg.textContent=''}, 3000);
        } catch(e) { msg.textContent = '❌ Error'; }
    }

    loadStatus();
    </script>
</body>
</html>
"""


@app.get("/", response_class=HTMLResponse)
async def landing_page():
    return LANDING_PAGE


@app.get("/healthz")
def healthz():
    return {"ok": True, "port": PORT, "model": MODEL}


@app.post("/create-checkout-session")
async def create_checkout_session(request: Request):
    if not STRIPE_SECRET_KEY:
        raise HTTPException(status_code=500, detail="Stripe not configured")

    form_data = await request.form()
    product = form_data.get("product", "sprint_50")

    prices = {
        "sprint_50": {"price": 19900, "name": "50-Reply Sprint"},
        "lead_intake": {"price": 0, "name": "Free Sample (5 replies)"},
    }

    item = prices.get(product, prices["sprint_50"])

    try:
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[
                {
                    "price_data": {
                        "currency": "usd",
                        "product_data": {"name": item["name"]},
                        "unit_amount": item["price"],
                    },
                    "quantity": 1,
                }
            ],
            mode="payment",
            success_url=f"{request.base_url}success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{request.base_url}cancel",
            metadata={"product": product},
        )
        return RedirectResponse(url=checkout_session.url, status_code=303)
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=500, detail=f"Stripe error: {str(e)}")


@app.get("/success")
async def success_page(session_id: str):
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            return HTMLResponse("""
            <!DOCTYPE html>
            <html><head><script src="https://cdn.tailwindcss.com"></script></head>
            <body class="min-h-screen flex items-center justify-center" style="background:linear-gradient(135deg,#0f172a,#1e293b);color:#f1f5f9">
                <div class="text-center">
                    <div class="text-6xl mb-4">✅</div>
                    <h1 class="text-4xl font-extrabold mb-4">Payment Successful!</h1>
                    <p class="text-xl text-slate-300 mb-8">Thank you! We'll contact you within 24 hours to get started.</p>
                    <a href="mailto:droxai25@outlook.com" class="inline-block bg-indigo-600 text-white font-bold py-3 px-8 rounded-lg hover:bg-indigo-700">Email Us Now</a>
                </div>
            </body></html>
            """)
    except Exception:
        pass
    return HTMLResponse(
        "<h1>Order Status Unknown</h1><p>Please contact us at droxai25@outlook.com</p>"
    )


@app.get("/cancel")
async def cancel_page():
    return HTMLResponse("""
    <!DOCTYPE html>
    <html><head><script src="https://cdn.tailwindcss.com"></script></head>
    <body class="min-h-screen flex items-center justify-center" style="background:linear-gradient(135deg,#0f172a,#1e293b);color:#f1f5f9">
        <div class="text-center">
            <div class="text-6xl mb-4">❌</div>
            <h1 class="text-4xl font-extrabold mb-4">Checkout Cancelled</h1>
            <p class="text-xl text-slate-300 mb-8">No worries! Ready when you are.</p>
            <a href="/" class="inline-block bg-indigo-600 text-white font-bold py-3 px-8 rounded-lg hover:bg-indigo-700">Go Back</a>
        </div>
    </body></html>
    """)


@app.post("/webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")

    if not STRIPE_WEBHOOK_SECRET or not sig_header:
        raise HTTPException(status_code=400, detail="Webhook not configured")

    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, STRIPE_WEBHOOK_SECRET
        )
    except (ValueError, stripe.error.SignatureVerificationError) as e:
        raise HTTPException(status_code=400, detail=f"Invalid webhook: {str(e)}")

    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        print(
            f"Payment received! Session: {session['id']}, Product: {session.get('metadata', {}).get('product')}"
        )

    return {"status": "success"}


@app.post("/lead-intake")
async def lead_intake(lead: Lead):
    job = Job(
        business_name=lead.business_name,
        signoff=lead.signoff,
        tone_style=lead.tone_style,
        persona=lead.persona,
        compliance_profile=lead.compliance_profile,
        private_email=lead.private_email,
        banned_words=["guarantee"],
        max_length_per_reply=400,
        reviews=lead.reviews_sample,
    )
    replies = await generate_replies(job)
    base = job_dirs(lead.business_name)
    pdf = os.path.join(base, "sample-5.pdf")
    docx = os.path.join(base, "sample-5.docx")
    build_pdf(pdf, lead.business_name, job.reviews, replies, job.signoff)
    build_docx(docx, lead.business_name, job.reviews, replies, job.signoff)
    return {"ok": True, "files": [pdf, docx]}


# ---------------------------------------------------------------------------
# Google OAuth endpoints
# ---------------------------------------------------------------------------

# In-memory state store (single-user desktop app — one state token at a time)
_oauth_state: Optional[str] = None


@app.get("/oauth/start")
async def oauth_start():
    """Redirect the user to Google's OAuth2 consent screen."""
    from google_api import GOOGLE_CLIENT_ID, get_authorization_url  # noqa: PLC0415

    if not GOOGLE_CLIENT_ID:
        raise HTTPException(
            status_code=500,
            detail="GOOGLE_CLIENT_ID is not configured. Set it in engine.env.",
        )

    global _oauth_state
    auth_url, state = get_authorization_url()
    _oauth_state = state
    return RedirectResponse(url=auth_url)


@app.get("/oauth/callback")
async def oauth_callback(code: str, state: str):
    """Handle Google OAuth2 redirect, exchange code for tokens."""
    global _oauth_state
    from google_api import exchange_code_for_tokens  # noqa: PLC0415

    if state != _oauth_state:
        raise HTTPException(status_code=400, detail="Invalid OAuth state parameter.")

    try:
        exchange_code_for_tokens(code=code, state=state)
    except Exception as exc:
        logger.error("OAuth token exchange failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"OAuth failed: {exc}")

    _oauth_state = None
    return HTMLResponse("""
    <!DOCTYPE html>
    <html><head><script src="https://cdn.tailwindcss.com"></script></head>
    <body class="min-h-screen flex items-center justify-center"
          style="background:linear-gradient(135deg,#0f172a,#1e293b);color:#f1f5f9">
        <div class="text-center">
            <div class="text-6xl mb-4">✅</div>
            <h1 class="text-4xl font-extrabold mb-4">Google Account Connected!</h1>
            <p class="text-xl text-slate-300 mb-8">
                Your credentials are stored securely. The orchestrator will now
                automatically check for new reviews and post replies every hour.
            </p>
            <a href="/" class="inline-block bg-indigo-600 text-white font-bold py-3 px-8 rounded-lg hover:bg-indigo-700">
                Go to Dashboard
            </a>
        </div>
    </body></html>
    """)


# ---------------------------------------------------------------------------
# Review management endpoints
# ---------------------------------------------------------------------------


@app.post("/reviews/sync")
async def reviews_sync():
    """Manually trigger a review sync (fetch + reply to new reviews)."""
    from monitor import run_sync_once  # noqa: PLC0415

    result = await run_sync_once()
    return result


@app.get("/reviews/status")
async def reviews_status():
    """Return current monitoring status and recent review statistics."""
    from monitor import get_monitor_status, get_all_reviews  # noqa: PLC0415
    from google_api import tokens_exist  # noqa: PLC0415

    status = get_monitor_status()
    status["google_connected"] = tokens_exist()
    status["recent_reviews"] = get_all_reviews()[:20]
    return status


# ---------------------------------------------------------------------------
# Browser automation endpoints
# ---------------------------------------------------------------------------


@app.post("/automation/setup")
async def automation_setup(creds: AutomationCredentials):
    """
    Save Google Business credentials (encrypted) and auto-start the service.

    The password is encrypted with Fernet before being written to config.json.
    It is NEVER logged or returned in plain text.
    """
    from browser_automation import save_credentials  # noqa: PLC0415
    from automation_service import start_service  # noqa: PLC0415

    save_credentials(creds.email, creds.password)
    result = start_service()
    return {
        "ok": True,
        "message": "Credentials saved and automation started.",
        "service": result,
    }


@app.post("/automation/start")
async def automation_start():
    """Start the browser automation background service."""
    from automation_service import start_service  # noqa: PLC0415

    return start_service()


@app.post("/automation/stop")
async def automation_stop():
    """Stop the browser automation background service."""
    from automation_service import stop_service  # noqa: PLC0415

    return stop_service()


@app.get("/automation/status")
async def automation_status():
    """Return automation service status."""
    from automation_service import get_status  # noqa: PLC0415
    from browser_automation import credentials_configured  # noqa: PLC0415

    status = get_status()
    status["credentials_configured"] = credentials_configured()
    return status


@app.post("/automation/run-now")
async def automation_run_now():
    """Immediately trigger one automation pass (outside of the scheduled interval)."""
    from automation_service import run_once  # noqa: PLC0415

    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, run_once)
    return result


@app.post("/automation/config")
async def automation_config_update(cfg: AutomationConfig):
    """Update automation configuration settings."""
    from browser_automation import load_config, save_config  # noqa: PLC0415

    current = load_config()
    auto_cfg = current.setdefault("automation", {})

    if cfg.enabled is not None:
        auto_cfg["enabled"] = cfg.enabled
    if cfg.interval_minutes is not None:
        if cfg.interval_minutes < 1 or cfg.interval_minutes > 1440:
            raise HTTPException(
                status_code=400, detail="interval_minutes must be between 1 and 1440"
            )
        auto_cfg["interval_minutes"] = cfg.interval_minutes
    if cfg.headless is not None:
        auto_cfg["headless"] = cfg.headless
    if cfg.business_name is not None:
        auto_cfg["business_name"] = cfg.business_name

    save_config(current)
    return {"ok": True, "automation": auto_cfg}


@app.get("/automation/logs")
async def automation_logs(lines: int = 100):
    """Return recent lines from automation.log."""
    from automation_service import get_recent_logs  # noqa: PLC0415

    return {"lines": get_recent_logs(lines)}
